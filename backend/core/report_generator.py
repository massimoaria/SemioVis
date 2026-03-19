"""PDF / DOCX / HTML report generation for SemioVis.

Generates structured semiotic analysis reports following the framework of
Kress, G., & van Leeuwen, T. (2006). Reading Images: The Grammar of Visual
Design (2nd ed.). Routledge.
"""

import base64
import io
import logging
import uuid
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT, TA_RIGHT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import cm, mm
from reportlab.platypus import (
    BaseDocTemplate,
    Frame,
    Image,
    NextPageTemplate,
    PageBreak,
    PageTemplate,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)

logger = logging.getLogger(__name__)

REPORT_DIR = Path("/tmp/semiovis_reports")

# ---------------------------------------------------------------------------
# Colour constants
# ---------------------------------------------------------------------------
BRAND_DARK = colors.HexColor("#1a1a2e")
BRAND_PRIMARY = colors.HexColor("#16213e")
BRAND_ACCENT = colors.HexColor("#0f3460")
BRAND_HIGHLIGHT = colors.HexColor("#e94560")
TABLE_HEADER_BG = colors.HexColor("#16213e")
TABLE_HEADER_FG = colors.white
TABLE_ALT_ROW = colors.HexColor("#f0f4f8")
SECTION_LINE_COLOR = colors.HexColor("#0f3460")

# ---------------------------------------------------------------------------
# Styles
# ---------------------------------------------------------------------------

def _build_styles() -> dict[str, ParagraphStyle]:
    """Build the paragraph styles used throughout the report."""
    base = getSampleStyleSheet()
    styles: dict[str, ParagraphStyle] = {}

    styles["title"] = ParagraphStyle(
        "ReportTitle",
        parent=base["Title"],
        fontSize=26,
        leading=32,
        alignment=TA_CENTER,
        textColor=BRAND_DARK,
        spaceAfter=6 * mm,
    )
    styles["subtitle"] = ParagraphStyle(
        "ReportSubtitle",
        parent=base["Normal"],
        fontSize=12,
        leading=16,
        alignment=TA_CENTER,
        textColor=colors.grey,
        spaceAfter=20 * mm,
    )
    styles["heading1"] = ParagraphStyle(
        "SectionHeading",
        parent=base["Heading1"],
        fontSize=18,
        leading=22,
        textColor=BRAND_DARK,
        spaceBefore=12 * mm,
        spaceAfter=4 * mm,
        borderWidth=0,
    )
    styles["heading2"] = ParagraphStyle(
        "SubHeading",
        parent=base["Heading2"],
        fontSize=14,
        leading=18,
        textColor=BRAND_ACCENT,
        spaceBefore=6 * mm,
        spaceAfter=3 * mm,
    )
    styles["body"] = ParagraphStyle(
        "BodyText",
        parent=base["Normal"],
        fontSize=10,
        leading=14,
        alignment=TA_JUSTIFY,
        spaceAfter=3 * mm,
    )
    styles["body_small"] = ParagraphStyle(
        "BodySmall",
        parent=base["Normal"],
        fontSize=9,
        leading=12,
        alignment=TA_LEFT,
        spaceAfter=2 * mm,
        textColor=colors.HexColor("#444444"),
    )
    styles["label"] = ParagraphStyle(
        "Label",
        parent=base["Normal"],
        fontSize=9,
        leading=11,
        textColor=colors.HexColor("#666666"),
    )
    styles["cell"] = ParagraphStyle(
        "CellText",
        parent=base["Normal"],
        fontSize=8,
        leading=10,
        alignment=TA_LEFT,
    )
    styles["cell_header"] = ParagraphStyle(
        "CellHeader",
        parent=base["Normal"],
        fontSize=8,
        leading=10,
        alignment=TA_LEFT,
        textColor=TABLE_HEADER_FG,
        fontName="Helvetica-Bold",
    )
    styles["interpretation"] = ParagraphStyle(
        "Interpretation",
        parent=base["Normal"],
        fontSize=10,
        leading=14,
        alignment=TA_JUSTIFY,
        leftIndent=8 * mm,
        rightIndent=8 * mm,
        spaceBefore=3 * mm,
        spaceAfter=4 * mm,
        backColor=colors.HexColor("#f8f9fa"),
        borderPadding=6,
    )
    styles["reference"] = ParagraphStyle(
        "Reference",
        parent=base["Normal"],
        fontSize=10,
        leading=14,
        alignment=TA_LEFT,
        leftIndent=12 * mm,
        spaceBefore=2 * mm,
    )
    styles["footer"] = ParagraphStyle(
        "Footer",
        parent=base["Normal"],
        fontSize=8,
        leading=10,
        alignment=TA_CENTER,
        textColor=colors.grey,
    )
    return styles


# ---------------------------------------------------------------------------
# Helper utilities
# ---------------------------------------------------------------------------

def _safe_get(data: dict | None, *keys: str, default: Any = "") -> Any:
    """Safely traverse nested dicts."""
    current = data or {}
    for k in keys:
        if isinstance(current, dict):
            current = current.get(k, default)
        else:
            return default
    return current


def _fmt(value: Any, precision: int = 2) -> str:
    """Format a numeric value for display."""
    if value is None:
        return "N/A"
    if isinstance(value, float):
        return f"{value:.{precision}f}"
    return str(value)


def _section_separator() -> Table:
    """Return a thin horizontal line as a section separator."""
    t = Table([[""]], colWidths=[16 * cm])
    t.setStyle(TableStyle([
        ("LINEBELOW", (0, 0), (-1, -1), 1, SECTION_LINE_COLOR),
    ]))
    return t


def _make_table(
    headers: list[str],
    rows: list[list[str]],
    col_widths: list[float] | None = None,
    styles_dict: dict | None = None,
) -> Table:
    """Create a styled table with header row and alternating row colours."""
    s = styles_dict or _build_styles()
    header_cells = [Paragraph(h, s["cell_header"]) for h in headers]
    body_rows = [
        [Paragraph(str(cell), s["cell"]) for cell in row] for row in rows
    ]
    data = [header_cells] + body_rows
    t = Table(data, colWidths=col_widths, repeatRows=1)

    style_commands: list[tuple] = [
        ("BACKGROUND", (0, 0), (-1, 0), TABLE_HEADER_BG),
        ("TEXTCOLOR", (0, 0), (-1, 0), TABLE_HEADER_FG),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("FONTSIZE", (0, 0), (-1, 0), 8),
        ("BOTTOMPADDING", (0, 0), (-1, 0), 6),
        ("TOPPADDING", (0, 0), (-1, 0), 6),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#cccccc")),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("LEFTPADDING", (0, 0), (-1, -1), 4),
        ("RIGHTPADDING", (0, 0), (-1, -1), 4),
    ]
    # Alternating row colours
    for i in range(1, len(data)):
        if i % 2 == 0:
            style_commands.append(("BACKGROUND", (0, i), (-1, i), TABLE_ALT_ROW))

    t.setStyle(TableStyle(style_commands))
    return t


def _thumbnail_image(base64_str: str | None, max_width: float = 8 * cm) -> Image | None:
    """Convert a base64-encoded image string to a reportlab Image flowable."""
    if not base64_str:
        return None
    try:
        img_data = base64.b64decode(base64_str)
        buf = io.BytesIO(img_data)
        img = Image(buf)
        # Scale to fit within max_width while keeping aspect ratio
        aspect = img.imageHeight / img.imageWidth if img.imageWidth else 1
        img.drawWidth = max_width
        img.drawHeight = max_width * aspect
        # Cap the height
        max_height = 10 * cm
        if img.drawHeight > max_height:
            img.drawHeight = max_height
            img.drawWidth = max_height / aspect
        img.hAlign = "CENTER"
        return img
    except Exception as exc:
        logger.warning("Failed to decode thumbnail for report: %s", exc)
        return None


def _colour_swatch_table(
    palette: list[dict], styles_dict: dict
) -> Table | None:
    """Build a small table showing colour swatches from the palette."""
    if not palette:
        return None

    header = ["Colour", "Hex", "Proportion", "Zone"]
    rows: list[list[Any]] = []
    for swatch in palette[:8]:
        hex_val = swatch.get("hex", "#000000")
        proportion = swatch.get("proportion", 0)
        zone = swatch.get("zone_association", "")
        # A small coloured rectangle represented as a tiny table cell
        rows.append([hex_val, hex_val, f"{proportion:.1%}", zone])

    data = [[Paragraph(h, styles_dict["cell_header"]) for h in header]]
    for row in rows:
        data.append([Paragraph(str(c), styles_dict["cell"]) for c in row])

    t = Table(data, colWidths=[3 * cm, 3 * cm, 3 * cm, 5 * cm], repeatRows=1)

    style_commands: list[tuple] = [
        ("BACKGROUND", (0, 0), (-1, 0), TABLE_HEADER_BG),
        ("TEXTCOLOR", (0, 0), (-1, 0), TABLE_HEADER_FG),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#cccccc")),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("LEFTPADDING", (0, 0), (-1, -1), 4),
        ("RIGHTPADDING", (0, 0), (-1, -1), 4),
        ("TOPPADDING", (0, 0), (-1, -1), 3),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
    ]
    # Colour the first column cells with the actual swatch colour
    for i, swatch in enumerate(palette[:8], start=1):
        hex_val = swatch.get("hex", "#000000")
        try:
            bg = colors.HexColor(hex_val)
            style_commands.append(("BACKGROUND", (0, i), (0, i), bg))
            # Use white text if background is dark
            r, g, b = bg.red, bg.green, bg.blue
            luminance = 0.299 * r + 0.587 * g + 0.114 * b
            fg = colors.white if luminance < 0.5 else colors.black
            style_commands.append(("TEXTCOLOR", (0, i), (0, i), fg))
        except Exception:
            pass

    for i in range(1, len(data)):
        if i % 2 == 0:
            style_commands.append(("BACKGROUND", (1, i), (-1, i), TABLE_ALT_ROW))

    t.setStyle(TableStyle(style_commands))
    return t


# ---------------------------------------------------------------------------
# PDF page decorators
# ---------------------------------------------------------------------------

def _title_page_decorator(canvas, doc):
    """Draw the title page background and footer."""
    canvas.saveState()
    # Top accent bar
    canvas.setFillColor(BRAND_DARK)
    canvas.rect(0, A4[1] - 2.5 * cm, A4[0], 2.5 * cm, fill=True, stroke=False)
    # Bottom accent bar
    canvas.setFillColor(BRAND_HIGHLIGHT)
    canvas.rect(0, 0, A4[0], 4 * mm, fill=True, stroke=False)
    canvas.restoreState()


def _body_page_decorator(canvas, doc):
    """Draw header line and page number on body pages."""
    canvas.saveState()
    # Header line
    canvas.setStrokeColor(BRAND_ACCENT)
    canvas.setLineWidth(0.5)
    canvas.line(2 * cm, A4[1] - 1.5 * cm, A4[0] - 2 * cm, A4[1] - 1.5 * cm)
    # Header text
    canvas.setFont("Helvetica", 8)
    canvas.setFillColor(colors.grey)
    canvas.drawString(2 * cm, A4[1] - 1.3 * cm, "SemioVis — Semiotic Image Analysis Report")
    # Page number
    canvas.drawRightString(
        A4[0] - 2 * cm, 1.2 * cm, f"Page {doc.page}"
    )
    # Bottom accent line
    canvas.setStrokeColor(BRAND_HIGHLIGHT)
    canvas.setLineWidth(1)
    canvas.line(2 * cm, 1 * cm, A4[0] - 2 * cm, 1 * cm)
    canvas.restoreState()


# ---------------------------------------------------------------------------
# PDF builder
# ---------------------------------------------------------------------------

def _build_pdf(image_id: str, results: dict, output_path: Path) -> Path:
    """Build a multi-page PDF report using reportlab."""
    styles = _build_styles()

    # Configure document with two page templates: title page and body pages
    title_frame = Frame(
        2 * cm, 2 * cm, A4[0] - 4 * cm, A4[1] - 5 * cm,
        id="title_frame",
    )
    body_frame = Frame(
        2 * cm, 2 * cm, A4[0] - 4 * cm, A4[1] - 4 * cm,
        id="body_frame",
    )

    doc = BaseDocTemplate(
        str(output_path),
        pagesize=A4,
        leftMargin=2 * cm,
        rightMargin=2 * cm,
        topMargin=2 * cm,
        bottomMargin=2 * cm,
        title="SemioVis — Semiotic Image Analysis Report",
        author="SemioVis",
    )
    doc.addPageTemplates([
        PageTemplate(id="title", frames=[title_frame], onPage=_title_page_decorator),
        PageTemplate(id="body", frames=[body_frame], onPage=_body_page_decorator),
    ])

    story: list = []

    # ------------------------------------------------------------------
    # 1. Title page
    # ------------------------------------------------------------------
    story.append(Spacer(1, 6 * cm))
    story.append(Paragraph(
        "SemioVis",
        ParagraphStyle("BigTitle", parent=styles["title"], fontSize=36, leading=42),
    ))
    story.append(Paragraph(
        "Semiotic Image Analysis Report",
        styles["title"],
    ))
    story.append(Spacer(1, 1 * cm))
    story.append(Paragraph(
        f"Image ID: {image_id}",
        styles["subtitle"],
    ))
    story.append(Paragraph(
        f"Generated: {datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M UTC')}",
        styles["subtitle"],
    ))

    # Thumbnail on title page
    thumbnail_b64 = results.get("thumbnail_base64")
    thumb = _thumbnail_image(thumbnail_b64, max_width=7 * cm)
    if thumb:
        story.append(Spacer(1, 1 * cm))
        story.append(thumb)

    story.append(NextPageTemplate("body"))
    story.append(PageBreak())

    # ------------------------------------------------------------------
    # 2. Image metadata
    # ------------------------------------------------------------------
    story.append(Paragraph("1. Image Metadata", styles["heading1"]))
    story.append(_section_separator())

    meta = results.get("image_meta") or {}
    meta_items = [
        ("Image ID", image_id),
        ("Width", _safe_get(meta, "width", default="N/A")),
        ("Height", _safe_get(meta, "height", default="N/A")),
        ("Format", _safe_get(meta, "format", default="N/A")),
    ]
    for label, value in meta_items:
        story.append(Paragraph(
            f"<b>{label}:</b> {value}", styles["body"]
        ))

    # Thumbnail in metadata section
    if thumb:
        story.append(Spacer(1, 4 * mm))
        thumb2 = _thumbnail_image(thumbnail_b64, max_width=10 * cm)
        if thumb2:
            story.append(thumb2)
        story.append(Spacer(1, 4 * mm))

    # ------------------------------------------------------------------
    # 3. Representational Analysis
    # ------------------------------------------------------------------
    rep = results.get("representational") or {}
    story.append(Paragraph("2. Representational Analysis", styles["heading1"]))
    story.append(_section_separator())

    structure_type = _safe_get(rep, "structure_type", default="N/A")
    subtype = (
        _safe_get(rep, "narrative_subtype")
        or _safe_get(rep, "conceptual_subtype")
        or "N/A"
    )
    story.append(Paragraph(
        f"<b>Structure type:</b> {structure_type}", styles["body"]
    ))
    story.append(Paragraph(
        f"<b>Subtype:</b> {subtype.replace('_', ' ')}", styles["body"]
    ))
    story.append(Paragraph(
        f"<b>Vector count:</b> {_safe_get(rep, 'vector_count', default=0)}", styles["body"]
    ))
    story.append(Paragraph(
        f"<b>Dominant direction:</b> {_safe_get(rep, 'dominant_direction', default='N/A')}",
        styles["body"],
    ))

    # Participants table
    participants = _safe_get(rep, "participants", default=[])
    if participants:
        story.append(Paragraph("Participants", styles["heading2"]))
        p_headers = ["Label", "Confidence", "Human", "Animal", "BBox"]
        p_rows = []
        for p in participants:
            bbox = p.get("bbox", [])
            bbox_str = ", ".join(_fmt(v, 3) for v in bbox) if bbox else "N/A"
            p_rows.append([
                p.get("label", ""),
                _fmt(p.get("confidence", 0)),
                "Yes" if p.get("is_human") else "No",
                "Yes" if p.get("is_animal") else "No",
                bbox_str,
            ])
        story.append(_make_table(
            p_headers, p_rows,
            col_widths=[3.5 * cm, 2.5 * cm, 2 * cm, 2 * cm, 4 * cm],
            styles_dict=styles,
        ))
        story.append(Spacer(1, 3 * mm))

    # Interpretation
    interp_text = _safe_get(rep, "interpretation", default="")
    if interp_text:
        story.append(Paragraph("Interpretation", styles["heading2"]))
        story.append(Paragraph(interp_text, styles["interpretation"]))

    # ------------------------------------------------------------------
    # 4. Interactive Analysis
    # ------------------------------------------------------------------
    inter = results.get("interactive") or {}
    story.append(Paragraph("3. Interactive Analysis", styles["heading1"]))
    story.append(_section_separator())

    story.append(Paragraph(
        f"<b>Vertical angle:</b> {_safe_get(inter, 'vertical_angle', default='N/A').replace('_', ' ')}",
        styles["body"],
    ))
    story.append(Paragraph(
        f"<b>Horizontal angle:</b> {_safe_get(inter, 'horizontal_angle', default='N/A')}",
        styles["body"],
    ))
    story.append(Paragraph(
        f"<b>Power relation:</b> {_safe_get(inter, 'power_relation', default='N/A').replace('_', ' ')}",
        styles["body"],
    ))
    story.append(Paragraph(
        f"<b>Involvement:</b> {_safe_get(inter, 'involvement', default='N/A')}",
        styles["body"],
    ))
    story.append(Paragraph(
        f"<b>Coding orientation:</b> {_safe_get(inter, 'coding_orientation', default='N/A')}",
        styles["body"],
    ))
    story.append(Paragraph(
        f"<b>Modality score:</b> {_fmt(_safe_get(inter, 'modality_score', default=None))}",
        styles["body"],
    ))

    vp = _safe_get(inter, "vanishing_point", default=None)
    if vp and isinstance(vp, (list, tuple)) and len(vp) == 2:
        story.append(Paragraph(
            f"<b>Vanishing point:</b> ({_fmt(vp[0])}, {_fmt(vp[1])})",
            styles["body"],
        ))

    # Faces table
    faces = _safe_get(inter, "faces", default=[])
    if faces:
        story.append(Paragraph("Face Analysis", styles["heading2"]))
        f_headers = ["ID", "Gaze", "Distance", "Shot", "Pan", "Tilt", "Top Emotion"]
        f_rows = []
        for face in faces:
            emotions = face.get("emotions", {})
            top_emotion = max(emotions, key=emotions.get) if emotions else "N/A"
            top_score = _fmt(emotions.get(top_emotion, 0)) if emotions else ""
            f_rows.append([
                str(face.get("face_id", "")),
                face.get("gaze_type", ""),
                face.get("social_distance", "").replace("_", " "),
                face.get("shot_type", "").replace("_", " "),
                _fmt(face.get("pan_angle", 0), 1) + "\u00b0",
                _fmt(face.get("tilt_angle", 0), 1) + "\u00b0",
                f"{top_emotion} ({top_score})",
            ])
        story.append(_make_table(
            f_headers, f_rows,
            col_widths=[1.2 * cm, 2 * cm, 2.5 * cm, 2.5 * cm, 1.8 * cm, 1.8 * cm, 3.5 * cm],
            styles_dict=styles,
        ))
        story.append(Spacer(1, 3 * mm))

    # Modality profile table (8 markers)
    modality = _safe_get(inter, "modality_profile", default={})
    if modality:
        story.append(Paragraph("Modality Profile (8 Markers)", styles["heading2"]))
        marker_names = [
            ("colour_saturation", "Colour Saturation"),
            ("colour_differentiation", "Colour Differentiation"),
            ("colour_modulation", "Colour Modulation"),
            ("contextualization", "Contextualization"),
            ("representation", "Representation"),
            ("depth", "Depth"),
            ("illumination", "Illumination"),
            ("brightness", "Brightness"),
        ]
        m_headers = ["Marker", "Value (0-1)", "Level"]
        m_rows = []
        for key, display_name in marker_names:
            val = modality.get(key)
            if val is not None:
                level = "Low" if val < 0.33 else ("Medium" if val < 0.66 else "High")
                m_rows.append([display_name, _fmt(val), level])
        if m_rows:
            story.append(_make_table(
                m_headers, m_rows,
                col_widths=[5.5 * cm, 3 * cm, 3 * cm],
                styles_dict=styles,
            ))
            story.append(Spacer(1, 3 * mm))

    # Interactive interpretation
    inter_interp = _safe_get(inter, "interpretation", default="")
    if inter_interp:
        story.append(Paragraph("Interpretation", styles["heading2"]))
        story.append(Paragraph(inter_interp, styles["interpretation"]))

    # ------------------------------------------------------------------
    # 5. Compositional Analysis
    # ------------------------------------------------------------------
    comp = results.get("compositional") or {}
    story.append(Paragraph("4. Compositional Analysis", styles["heading1"]))
    story.append(_section_separator())

    comp_type = _safe_get(comp, "composition_type", default="N/A")
    story.append(Paragraph(
        f"<b>Composition type:</b> {comp_type}", styles["body"]
    ))

    if comp_type == "centred":
        centred_sub = _safe_get(comp, "centred_subtype", default="N/A")
        story.append(Paragraph(
            f"<b>Centred subtype:</b> {centred_sub.replace('_', ' ') if centred_sub else 'N/A'}",
            styles["body"],
        ))
    else:
        axes = _safe_get(comp, "polarization_axes", default=[])
        if axes:
            axes_str = ", ".join(a.replace("_", "/") for a in axes)
            story.append(Paragraph(
                f"<b>Polarization axes:</b> {axes_str}", styles["body"]
            ))

    has_triptych = _safe_get(comp, "has_triptych", default=False)
    if has_triptych:
        orientation = _safe_get(comp, "triptych_orientation", default="")
        story.append(Paragraph(
            f"<b>Triptych:</b> Yes ({orientation})", styles["body"]
        ))

    dom_structure = _safe_get(comp, "dominant_structure", default="N/A")
    story.append(Paragraph(
        f"<b>Dominant structure:</b> {dom_structure.replace('_', ' ')}",
        styles["body"],
    ))

    # Zones table
    zones = _safe_get(comp, "zones", default=[])
    if zones:
        story.append(Paragraph("Spatial Zones", styles["heading2"]))
        z_headers = ["Zone", "Position", "Semiotic Label", "Saliency", "Weight", "Info Value"]
        z_rows = []
        for z in zones:
            z_rows.append([
                z.get("zone_id", ""),
                z.get("position_label", ""),
                z.get("semiotic_label", ""),
                _fmt(z.get("mean_saliency", 0)),
                _fmt(z.get("visual_weight", 0)),
                _fmt(z.get("information_value_score", 0)),
            ])
        story.append(_make_table(
            z_headers, z_rows,
            col_widths=[1.8 * cm, 2.5 * cm, 3 * cm, 2.2 * cm, 2.2 * cm, 2.5 * cm],
            styles_dict=styles,
        ))
        story.append(Spacer(1, 3 * mm))

    # Colour palette
    palette = _safe_get(comp, "color_palette", default=[])
    if palette:
        story.append(Paragraph("Colour Palette", styles["heading2"]))
        palette_table = _colour_swatch_table(palette, styles)
        if palette_table:
            story.append(palette_table)
            story.append(Spacer(1, 3 * mm))

    # Framing analysis
    framing = _safe_get(comp, "framing", default={})
    if framing:
        story.append(Paragraph("Framing Analysis", styles["heading2"]))
        disc = _safe_get(framing, "disconnection_score", default=None)
        conn = _safe_get(framing, "connection_score", default=None)
        story.append(Paragraph(
            f"<b>Disconnection score:</b> {_fmt(disc)}", styles["body"]
        ))
        story.append(Paragraph(
            f"<b>Connection score:</b> {_fmt(conn)}", styles["body"]
        ))
        story.append(Paragraph(
            f"<b>Frame lines:</b> {len(framing.get('frame_lines', []))} detected",
            styles["body"],
        ))
        story.append(Paragraph(
            f"<b>Empty space regions:</b> {framing.get('empty_space_regions', 0)}"
            f" &nbsp;|&nbsp; <b>Colour discontinuities:</b> {framing.get('colour_discontinuities', 0)}"
            f" &nbsp;|&nbsp; <b>Colour continuities:</b> {framing.get('colour_continuities', 0)}",
            styles["body"],
        ))
        story.append(Paragraph(
            f"<b>Visual vectors:</b> {framing.get('visual_vectors', 0)}"
            f" &nbsp;|&nbsp; <b>Shape rhymes:</b> {framing.get('shape_rhymes', 0)}",
            styles["body"],
        ))

    # Reading path
    reading_path = _safe_get(comp, "reading_path", default={})
    if reading_path:
        story.append(Paragraph("Reading Path", styles["heading2"]))
        path_shape = _safe_get(reading_path, "path_shape", default="N/A")
        is_linear = _safe_get(reading_path, "is_linear", default=False)
        waypoints = _safe_get(reading_path, "waypoints", default=[])
        story.append(Paragraph(
            f"<b>Path shape:</b> {path_shape.replace('_', ' ')}", styles["body"]
        ))
        story.append(Paragraph(
            f"<b>Linear:</b> {'Yes' if is_linear else 'No'}", styles["body"]
        ))
        story.append(Paragraph(
            f"<b>Waypoints:</b> {len(waypoints)}", styles["body"]
        ))

        if waypoints:
            wp_headers = ["#", "X", "Y", "Saliency", "Label"]
            wp_rows = []
            for i, wp in enumerate(waypoints[:10], start=1):
                wp_rows.append([
                    str(i),
                    _fmt(wp.get("x", 0), 3),
                    _fmt(wp.get("y", 0), 3),
                    _fmt(wp.get("saliency", 0)),
                    wp.get("label", ""),
                ])
            story.append(_make_table(
                wp_headers, wp_rows,
                col_widths=[1.2 * cm, 2.5 * cm, 2.5 * cm, 2.5 * cm, 5 * cm],
                styles_dict=styles,
            ))
            story.append(Spacer(1, 3 * mm))

    # Compositional interpretation
    comp_interp = _safe_get(comp, "interpretation", default="")
    if comp_interp:
        story.append(Paragraph("Interpretation", styles["heading2"]))
        story.append(Paragraph(comp_interp, styles["interpretation"]))

    # ------------------------------------------------------------------
    # 6. Full Interpretation
    # ------------------------------------------------------------------
    full_interp = results.get("full_interpretation") or results.get("interpretation")
    if full_interp:
        story.append(Paragraph("5. Integrated Semiotic Interpretation", styles["heading1"]))
        story.append(_section_separator())
        story.append(Paragraph(full_interp, styles["interpretation"]))

    # ------------------------------------------------------------------
    # 7. Reference
    # ------------------------------------------------------------------
    story.append(Paragraph("References", styles["heading1"]))
    story.append(_section_separator())
    story.append(Paragraph(
        "Kress, G., &amp; van Leeuwen, T. (2006). <i>Reading Images: The Grammar of "
        "Visual Design</i> (2nd ed.). Routledge.",
        styles["reference"],
    ))
    story.append(Spacer(1, 4 * mm))
    story.append(Paragraph(
        "Aria, M. (2026). SemioVis: Visual Social Semiotics Analysis Tool. "
        '<a href="https://github.com/massimoaria/SemioVis">'
        "https://github.com/massimoaria/SemioVis</a>",
        styles["reference"],
    ))
    story.append(Spacer(1, 8 * mm))
    story.append(Paragraph(
        "Report generated by SemioVis — https://github.com/massimoaria/SemioVis",
        styles["footer"],
    ))

    # Build the PDF
    doc.build(story)
    return output_path


# ---------------------------------------------------------------------------
# DOCX builder
# ---------------------------------------------------------------------------

def _build_docx(image_id: str, results: dict, output_path: Path) -> Path:
    """Build a DOCX report using python-docx."""
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt, RGBColor

    doc = Document()

    # Style configuration
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(10)

    # ------------------------------------------------------------------
    # 1. Title
    # ------------------------------------------------------------------
    title_para = doc.add_heading("SemioVis", level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = doc.add_paragraph("Semiotic Image Analysis Report")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(14)
    subtitle.runs[0].font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph(f"Image ID: {image_id}").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(
        f"Generated: {datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M UTC')}"
    ).alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Thumbnail
    thumbnail_b64 = results.get("thumbnail_base64")
    if thumbnail_b64:
        try:
            img_data = base64.b64decode(thumbnail_b64)
            buf = io.BytesIO(img_data)
            doc.add_picture(buf, width=Inches(3))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        except Exception as exc:
            logger.warning("Failed to insert thumbnail into DOCX: %s", exc)

    doc.add_page_break()

    # ------------------------------------------------------------------
    # 2. Image Metadata
    # ------------------------------------------------------------------
    doc.add_heading("1. Image Metadata", level=1)
    meta = results.get("image_meta") or {}
    doc.add_paragraph(f"Image ID: {image_id}")
    for key in ("width", "height", "format"):
        doc.add_paragraph(f"{key.capitalize()}: {meta.get(key, 'N/A')}")

    # ------------------------------------------------------------------
    # 3. Representational Analysis
    # ------------------------------------------------------------------
    rep = results.get("representational") or {}
    doc.add_heading("2. Representational Analysis", level=1)
    doc.add_paragraph(f"Structure type: {_safe_get(rep, 'structure_type', default='N/A')}")
    subtype = _safe_get(rep, "narrative_subtype") or _safe_get(rep, "conceptual_subtype") or "N/A"
    doc.add_paragraph(f"Subtype: {subtype.replace('_', ' ')}")
    doc.add_paragraph(f"Vector count: {_safe_get(rep, 'vector_count', default=0)}")
    doc.add_paragraph(f"Dominant direction: {_safe_get(rep, 'dominant_direction', default='N/A')}")

    participants = _safe_get(rep, "participants", default=[])
    if participants:
        doc.add_heading("Participants", level=2)
        table = doc.add_table(rows=1, cols=4)
        table.style = "Light Grid Accent 1"
        hdr = table.rows[0].cells
        for i, h in enumerate(["Label", "Confidence", "Human", "Animal"]):
            hdr[i].text = h
        for p in participants:
            row = table.add_row().cells
            row[0].text = p.get("label", "")
            row[1].text = _fmt(p.get("confidence", 0))
            row[2].text = "Yes" if p.get("is_human") else "No"
            row[3].text = "Yes" if p.get("is_animal") else "No"

    interp_text = _safe_get(rep, "interpretation", default="")
    if interp_text:
        doc.add_heading("Interpretation", level=2)
        doc.add_paragraph(interp_text)

    # ------------------------------------------------------------------
    # 4. Interactive Analysis
    # ------------------------------------------------------------------
    inter = results.get("interactive") or {}
    doc.add_heading("3. Interactive Analysis", level=1)
    doc.add_paragraph(f"Vertical angle: {_safe_get(inter, 'vertical_angle', default='N/A')}")
    doc.add_paragraph(f"Horizontal angle: {_safe_get(inter, 'horizontal_angle', default='N/A')}")
    doc.add_paragraph(f"Power relation: {_safe_get(inter, 'power_relation', default='N/A')}")
    doc.add_paragraph(f"Involvement: {_safe_get(inter, 'involvement', default='N/A')}")
    doc.add_paragraph(f"Coding orientation: {_safe_get(inter, 'coding_orientation', default='N/A')}")
    doc.add_paragraph(f"Modality score: {_fmt(_safe_get(inter, 'modality_score', default=None))}")

    faces = _safe_get(inter, "faces", default=[])
    if faces:
        doc.add_heading("Face Analysis", level=2)
        table = doc.add_table(rows=1, cols=5)
        table.style = "Light Grid Accent 1"
        hdr = table.rows[0].cells
        for i, h in enumerate(["ID", "Gaze", "Distance", "Shot", "Top Emotion"]):
            hdr[i].text = h
        for face in faces:
            emotions = face.get("emotions", {})
            top_emotion = max(emotions, key=emotions.get) if emotions else "N/A"
            row = table.add_row().cells
            row[0].text = str(face.get("face_id", ""))
            row[1].text = face.get("gaze_type", "")
            row[2].text = face.get("social_distance", "").replace("_", " ")
            row[3].text = face.get("shot_type", "").replace("_", " ")
            row[4].text = top_emotion

    modality = _safe_get(inter, "modality_profile", default={})
    if modality:
        doc.add_heading("Modality Profile", level=2)
        table = doc.add_table(rows=1, cols=3)
        table.style = "Light Grid Accent 1"
        hdr = table.rows[0].cells
        for i, h in enumerate(["Marker", "Value", "Level"]):
            hdr[i].text = h
        for key, display_name in [
            ("colour_saturation", "Colour Saturation"),
            ("colour_differentiation", "Colour Differentiation"),
            ("colour_modulation", "Colour Modulation"),
            ("contextualization", "Contextualization"),
            ("representation", "Representation"),
            ("depth", "Depth"),
            ("illumination", "Illumination"),
            ("brightness", "Brightness"),
        ]:
            val = modality.get(key)
            if val is not None:
                level = "Low" if val < 0.33 else ("Medium" if val < 0.66 else "High")
                row = table.add_row().cells
                row[0].text = display_name
                row[1].text = _fmt(val)
                row[2].text = level

    inter_interp = _safe_get(inter, "interpretation", default="")
    if inter_interp:
        doc.add_heading("Interpretation", level=2)
        doc.add_paragraph(inter_interp)

    # ------------------------------------------------------------------
    # 5. Compositional Analysis
    # ------------------------------------------------------------------
    comp = results.get("compositional") or {}
    doc.add_heading("4. Compositional Analysis", level=1)
    doc.add_paragraph(f"Composition type: {_safe_get(comp, 'composition_type', default='N/A')}")
    doc.add_paragraph(f"Dominant structure: {_safe_get(comp, 'dominant_structure', default='N/A')}")

    zones = _safe_get(comp, "zones", default=[])
    if zones:
        doc.add_heading("Spatial Zones", level=2)
        table = doc.add_table(rows=1, cols=5)
        table.style = "Light Grid Accent 1"
        hdr = table.rows[0].cells
        for i, h in enumerate(["Zone", "Position", "Semiotic Label", "Saliency", "Weight"]):
            hdr[i].text = h
        for z in zones:
            row = table.add_row().cells
            row[0].text = z.get("zone_id", "")
            row[1].text = z.get("position_label", "")
            row[2].text = z.get("semiotic_label", "")
            row[3].text = _fmt(z.get("mean_saliency", 0))
            row[4].text = _fmt(z.get("visual_weight", 0))

    framing = _safe_get(comp, "framing", default={})
    if framing:
        doc.add_heading("Framing Analysis", level=2)
        doc.add_paragraph(
            f"Disconnection: {_fmt(framing.get('disconnection_score'))} | "
            f"Connection: {_fmt(framing.get('connection_score'))}"
        )

    reading_path = _safe_get(comp, "reading_path", default={})
    if reading_path:
        doc.add_heading("Reading Path", level=2)
        doc.add_paragraph(
            f"Shape: {_safe_get(reading_path, 'path_shape', default='N/A').replace('_', ' ')} | "
            f"Linear: {'Yes' if reading_path.get('is_linear') else 'No'} | "
            f"Waypoints: {len(reading_path.get('waypoints', []))}"
        )

    comp_interp = _safe_get(comp, "interpretation", default="")
    if comp_interp:
        doc.add_heading("Interpretation", level=2)
        doc.add_paragraph(comp_interp)

    # ------------------------------------------------------------------
    # 6. Full Interpretation
    # ------------------------------------------------------------------
    full_interp = results.get("full_interpretation") or results.get("interpretation")
    if full_interp:
        doc.add_heading("5. Integrated Semiotic Interpretation", level=1)
        doc.add_paragraph(full_interp)

    # ------------------------------------------------------------------
    # 7. Reference
    # ------------------------------------------------------------------
    doc.add_heading("References", level=1)
    doc.add_paragraph(
        "Kress, G., & van Leeuwen, T. (2006). Reading Images: The Grammar of "
        "Visual Design (2nd ed.). Routledge."
    )
    doc.add_paragraph(
        "Aria, M. (2026). SemioVis: Visual Social Semiotics Analysis Tool. "
        "https://github.com/massimoaria/SemioVis"
    )

    para = doc.add_paragraph("Report generated by SemioVis — https://github.com/massimoaria/SemioVis")
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.save(str(output_path))
    return output_path


# ---------------------------------------------------------------------------
# HTML builder
# ---------------------------------------------------------------------------

_HTML_TEMPLATE = """\
<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>SemioVis Report — {{ image_id }}</title>
<style>
  body { font-family: 'Segoe UI', Roboto, Helvetica, Arial, sans-serif;
         max-width: 900px; margin: 0 auto; padding: 2em; color: #222; }
  h1 { color: #1a1a2e; border-bottom: 2px solid #0f3460; padding-bottom: 0.3em; }
  h2 { color: #0f3460; }
  h3 { color: #16213e; }
  table { border-collapse: collapse; width: 100%; margin: 1em 0; }
  th { background: #16213e; color: #fff; padding: 8px 12px; text-align: left; font-size: 0.9em; }
  td { border: 1px solid #ddd; padding: 6px 12px; font-size: 0.9em; }
  tr:nth-child(even) { background: #f0f4f8; }
  .interpretation { background: #f8f9fa; padding: 1em 1.5em; border-left: 4px solid #0f3460;
                     margin: 1em 0; font-style: italic; }
  .title-block { text-align: center; margin: 3em 0; }
  .title-block h1 { font-size: 2.2em; border: none; }
  .title-block p { color: #666; }
  .thumbnail { text-align: center; margin: 1.5em 0; }
  .thumbnail img { max-width: 400px; border: 1px solid #ddd; border-radius: 4px; }
  .swatch { display: inline-block; width: 40px; height: 24px; border: 1px solid #999;
             border-radius: 3px; vertical-align: middle; margin-right: 6px; }
  .ref { font-style: italic; margin: 2em 0; }
  .footer { text-align: center; color: #999; font-size: 0.85em; margin-top: 3em;
             border-top: 1px solid #ddd; padding-top: 1em; }
  @media print { body { max-width: none; } }
</style>
</head>
<body>

<div class="title-block">
  <h1>SemioVis</h1>
  <p style="font-size:1.2em;">Semiotic Image Analysis Report</p>
  <p>Image ID: {{ image_id }}</p>
  <p>Generated: {{ generated_at }}</p>
</div>

{% if thumbnail_base64 %}
<div class="thumbnail">
  <img src="data:image/jpeg;base64,{{ thumbnail_base64 }}" alt="Analysed image">
</div>
{% endif %}

<h1>1. Image Metadata</h1>
<p><strong>Image ID:</strong> {{ image_id }}</p>
{% if meta %}
<p><strong>Width:</strong> {{ meta.get('width', 'N/A') }}
   | <strong>Height:</strong> {{ meta.get('height', 'N/A') }}
   | <strong>Format:</strong> {{ meta.get('format', 'N/A') }}</p>
{% endif %}

<h1>2. Representational Analysis</h1>
{% if rep %}
<p><strong>Structure type:</strong> {{ rep.get('structure_type', 'N/A') }}</p>
<p><strong>Subtype:</strong> {{ (rep.get('narrative_subtype') or rep.get('conceptual_subtype') or 'N/A')|replace('_', ' ') }}</p>
<p><strong>Vector count:</strong> {{ rep.get('vector_count', 0) }}
   | <strong>Dominant direction:</strong> {{ rep.get('dominant_direction', 'N/A') }}</p>

{% if rep.get('participants') %}
<h2>Participants</h2>
<table>
<tr><th>Label</th><th>Confidence</th><th>Human</th><th>Animal</th></tr>
{% for p in rep.participants %}
<tr><td>{{ p.get('label', '') }}</td>
    <td>{{ "%.2f"|format(p.get('confidence', 0)) }}</td>
    <td>{{ "Yes" if p.get('is_human') else "No" }}</td>
    <td>{{ "Yes" if p.get('is_animal') else "No" }}</td></tr>
{% endfor %}
</table>
{% endif %}

{% if rep.get('interpretation') %}
<h2>Interpretation</h2>
<div class="interpretation">{{ rep.interpretation }}</div>
{% endif %}
{% endif %}

<h1>3. Interactive Analysis</h1>
{% if inter %}
<p><strong>Vertical angle:</strong> {{ inter.get('vertical_angle', 'N/A')|replace('_', ' ') }}
   | <strong>Horizontal angle:</strong> {{ inter.get('horizontal_angle', 'N/A') }}</p>
<p><strong>Power relation:</strong> {{ inter.get('power_relation', 'N/A')|replace('_', ' ') }}
   | <strong>Involvement:</strong> {{ inter.get('involvement', 'N/A') }}</p>
<p><strong>Coding orientation:</strong> {{ inter.get('coding_orientation', 'N/A') }}
   | <strong>Modality score:</strong> {{ "%.2f"|format(inter.get('modality_score', 0)) }}</p>

{% if inter.get('faces') %}
<h2>Face Analysis</h2>
<table>
<tr><th>ID</th><th>Gaze</th><th>Distance</th><th>Shot</th><th>Pan</th><th>Tilt</th></tr>
{% for f in inter.faces %}
<tr><td>{{ f.face_id }}</td>
    <td>{{ f.gaze_type }}</td>
    <td>{{ f.get('social_distance', '')|replace('_', ' ') }}</td>
    <td>{{ f.get('shot_type', '')|replace('_', ' ') }}</td>
    <td>{{ "%.1f"|format(f.get('pan_angle', 0)) }}&deg;</td>
    <td>{{ "%.1f"|format(f.get('tilt_angle', 0)) }}&deg;</td></tr>
{% endfor %}
</table>
{% endif %}

{% if inter.get('modality_profile') %}
<h2>Modality Profile</h2>
<table>
<tr><th>Marker</th><th>Value</th><th>Level</th></tr>
{% for key, name in [('colour_saturation','Colour Saturation'),('colour_differentiation','Colour Differentiation'),('colour_modulation','Colour Modulation'),('contextualization','Contextualization'),('representation','Representation'),('depth','Depth'),('illumination','Illumination'),('brightness','Brightness')] %}
{% set val = inter.modality_profile.get(key, 0) %}
<tr><td>{{ name }}</td>
    <td>{{ "%.2f"|format(val) }}</td>
    <td>{{ "Low" if val < 0.33 else ("Medium" if val < 0.66 else "High") }}</td></tr>
{% endfor %}
</table>
{% endif %}

{% if inter.get('interpretation') %}
<h2>Interpretation</h2>
<div class="interpretation">{{ inter.interpretation }}</div>
{% endif %}
{% endif %}

<h1>4. Compositional Analysis</h1>
{% if comp %}
<p><strong>Composition type:</strong> {{ comp.get('composition_type', 'N/A') }}
   | <strong>Dominant structure:</strong> {{ comp.get('dominant_structure', 'N/A')|replace('_', ' ') }}</p>

{% if comp.get('zones') %}
<h2>Spatial Zones</h2>
<table>
<tr><th>Zone</th><th>Position</th><th>Semiotic Label</th><th>Saliency</th><th>Weight</th></tr>
{% for z in comp.zones %}
<tr><td>{{ z.zone_id }}</td>
    <td>{{ z.position_label }}</td>
    <td>{{ z.semiotic_label }}</td>
    <td>{{ "%.2f"|format(z.get('mean_saliency', 0)) }}</td>
    <td>{{ "%.2f"|format(z.get('visual_weight', 0)) }}</td></tr>
{% endfor %}
</table>
{% endif %}

{% if comp.get('color_palette') %}
<h2>Colour Palette</h2>
<table>
<tr><th>Colour</th><th>Hex</th><th>Proportion</th><th>Zone</th></tr>
{% for s in comp.color_palette %}
<tr><td><span class="swatch" style="background:{{ s.hex }};"></span></td>
    <td>{{ s.hex }}</td>
    <td>{{ "%.1f%%"|format(s.get('proportion', 0) * 100) }}</td>
    <td>{{ s.get('zone_association', '') }}</td></tr>
{% endfor %}
</table>
{% endif %}

{% if comp.get('framing') %}
<h2>Framing</h2>
<p><strong>Disconnection:</strong> {{ "%.2f"|format(comp.framing.get('disconnection_score', 0)) }}
   | <strong>Connection:</strong> {{ "%.2f"|format(comp.framing.get('connection_score', 0)) }}</p>
{% endif %}

{% if comp.get('reading_path') %}
<h2>Reading Path</h2>
<p><strong>Shape:</strong> {{ comp.reading_path.get('path_shape', 'N/A')|replace('_', ' ') }}
   | <strong>Linear:</strong> {{ "Yes" if comp.reading_path.get('is_linear') else "No" }}
   | <strong>Waypoints:</strong> {{ comp.reading_path.get('waypoints', [])|length }}</p>
{% endif %}

{% if comp.get('interpretation') %}
<h2>Interpretation</h2>
<div class="interpretation">{{ comp.interpretation }}</div>
{% endif %}
{% endif %}

{% if full_interpretation %}
<h1>5. Integrated Semiotic Interpretation</h1>
<div class="interpretation">{{ full_interpretation }}</div>
{% endif %}

<h1>References</h1>
<p class="ref">Kress, G., &amp; van Leeuwen, T. (2006). <em>Reading Images: The Grammar of
Visual Design</em> (2nd ed.). Routledge.</p>
<p class="ref">Aria, M. (2026). SemioVis: Visual Social Semiotics Analysis Tool.
<a href="https://github.com/massimoaria/SemioVis">https://github.com/massimoaria/SemioVis</a></p>

<div class="footer">Report generated by SemioVis &mdash; <a href="https://github.com/massimoaria/SemioVis">https://github.com/massimoaria/SemioVis</a></div>
</body>
</html>
"""


def _build_html(image_id: str, results: dict, output_path: Path) -> Path:
    """Build an HTML report using Jinja2."""
    from jinja2 import Environment

    env = Environment(autoescape=False)
    template = env.from_string(_HTML_TEMPLATE)

    html_content = template.render(
        image_id=image_id,
        generated_at=datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M UTC"),
        thumbnail_base64=results.get("thumbnail_base64"),
        meta=results.get("image_meta") or {},
        rep=results.get("representational") or {},
        inter=results.get("interactive") or {},
        comp=results.get("compositional") or {},
        full_interpretation=(
            results.get("full_interpretation") or results.get("interpretation") or ""
        ),
    )

    output_path.write_text(html_content, encoding="utf-8")
    return output_path


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------

async def generate_report(
    image_id: str,
    results: dict,
    format: str = "pdf",
) -> Path:
    """Generate an analysis report in the requested format.

    Args:
        image_id: Unique identifier of the analysed image.
        results: Dictionary containing analysis results. Expected keys:
            - "representational": dict matching RepresentationalResult fields
            - "interactive": dict matching InteractiveResult fields
            - "compositional": dict matching CompositionalResult fields
            - "thumbnail_base64": optional base64-encoded image thumbnail
            - "image_meta": optional dict with width, height, format
            - "full_interpretation": optional integrated interpretation text
        format: Output format — "pdf", "docx", or "html".

    Returns:
        Path to the generated report file in /tmp/semiovis_reports/.

    Raises:
        ValueError: If format is not one of "pdf", "docx", "html".
    """
    REPORT_DIR.mkdir(parents=True, exist_ok=True)

    report_id = uuid.uuid4().hex[:12]
    timestamp = datetime.now(timezone.utc).strftime("%Y%m%d_%H%M%S")

    match format:
        case "pdf":
            filename = f"semiovis_report_{timestamp}_{report_id}.pdf"
            output_path = REPORT_DIR / filename
            return _build_pdf(image_id, results, output_path)
        case "docx":
            filename = f"semiovis_report_{timestamp}_{report_id}.docx"
            output_path = REPORT_DIR / filename
            return _build_docx(image_id, results, output_path)
        case "html":
            filename = f"semiovis_report_{timestamp}_{report_id}.html"
            output_path = REPORT_DIR / filename
            return _build_html(image_id, results, output_path)
        case _:
            raise ValueError(
                f"Unsupported report format: {format!r}. "
                f"Supported formats: pdf, docx, html."
            )
